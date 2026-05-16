"""Render Note rows to PDF and DOCX.

Both formats start from `note.note_content_markdown` (the Ollama-formatted
output that's also rendered on screen). We do a small, intentionally limited
markdown parse — headings, paragraphs, lists, horizontal rules, plus inline
bold/italic/inline-code — because that covers everything our templates can
produce. Avoiding a full markdown→HTML→PDF pipeline keeps the dependency
footprint to pure-Python wheels (no Cairo / Pango / pandoc system deps).
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import TYPE_CHECKING, Optional

from docx import Document
from reportlab.lib.pagesizes import letter

from app.services.diarization import relabel_speakers
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

if TYPE_CHECKING:
    from app.models import Note


@dataclass
class _Block:
    kind: str  # 'h1'..'h4', 'p', 'bullet', 'ordered', 'hr'
    text: str


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_HR_RE = re.compile(r"^[-*_]{3,}\s*$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.+)$")
_ORDERED_RE = re.compile(r"^\s*\d+\.\s+(.+)$")
_BREAKS_PARA_RE = re.compile(r"^(#{1,6}\s|[-*_]{3,}\s*$|\s*[-*]\s|\s*\d+\.\s)")


def _parse_blocks(md: str) -> list[_Block]:
    """Line-oriented markdown → flat block list.

    Consecutive non-blank, non-special lines fold into a single paragraph.
    Inline formatting stays as-is in `text`; renderers parse it per-format.
    """
    if not md:
        return []

    out: list[_Block] = []
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            out.append(_Block(kind=f"h{level}", text=m.group(2)))
            i += 1
            continue

        if _HR_RE.match(line):
            out.append(_Block(kind="hr", text=""))
            i += 1
            continue

        m = _BULLET_RE.match(line)
        if m:
            out.append(_Block(kind="bullet", text=m.group(1)))
            i += 1
            continue

        m = _ORDERED_RE.match(line)
        if m:
            out.append(_Block(kind="ordered", text=m.group(1)))
            i += 1
            continue

        # Plain paragraph — gather following lines that aren't blank or
        # block-starting markers.
        para = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt or _BREAKS_PARA_RE.match(nxt):
                break
            para.append(nxt)
            i += 1
        out.append(_Block(kind="p", text=" ".join(para)))

    return out


# Inline parser: handles **bold**, *italic*, `code`. Stops short of full
# markdown (links, images, nested emphasis) — those rarely appear in
# clinical-note output.
_INLINE_RE = re.compile(r"(\*\*[^*\n]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*)|`[^`\n]+`)")


def _html_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_to_rl(text: str) -> str:
    """Inline markdown → reportlab Paragraph mini-HTML."""
    s = _html_escape(text)
    s = re.sub(r"\*\*([^*\n]+)\*\*", r"<b>\1</b>", s)
    s = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"<i>\1</i>", s)
    s = re.sub(r"`([^`\n]+)`", r'<font face="Courier">\1</font>', s)
    return s


def _add_inline_runs_docx(paragraph, text: str) -> None:
    """Walk inline markdown, append styled runs to a python-docx paragraph."""
    cursor = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > cursor:
            paragraph.add_run(text[cursor:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        else:  # italic
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        cursor = m.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _participant_names(note: "Note") -> list[str]:
    names = []
    for p in note.participants:
        full = " ".join(filter(None, [p.first_name, p.last_name])).strip()
        if full:
            names.append(full)
    return names


def _doc_title(note: "Note", template_name: Optional[str]) -> str:
    if template_name:
        return template_name
    if note.note_date:
        return f"Note — {note.note_date.strftime('%B %d, %Y')}"
    return "Note"


# --- PDF ---------------------------------------------------------------------

def render_pdf(note: "Note", *, template_name: Optional[str] = None) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=_doc_title(note, template_name),
        author=note.author_name or "",
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    body = styles["BodyText"]
    body.spaceAfter = 6
    head = {1: styles["Heading1"], 2: styles["Heading2"], 3: styles["Heading3"], 4: styles["Heading4"]}

    story = [Paragraph(_html_escape(_doc_title(note, template_name)), title_style)]

    meta = []
    if note.note_date:
        meta.append(f"<b>Date:</b> {note.note_date.strftime('%B %d, %Y')}")
    if note.author_name:
        meta.append(f"<b>Author:</b> {_html_escape(note.author_name)}")
    if template_name:
        meta.append(f"<b>Template:</b> {_html_escape(template_name)}")
    participants = _participant_names(note)
    if participants:
        meta.append("<b>Participants:</b> " + _html_escape(", ".join(participants)))
    for line in meta:
        story.append(Paragraph(line, body))

    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color="#000000"))
    story.append(Spacer(1, 12))

    # Rewrite anonymous "Speaker N" labels to the named participants the user
    # assigned, so the exported document reads in real names.
    markdown = relabel_speakers(note.note_content_markdown or "", note.speaker_labels)
    for blk in _parse_blocks(markdown or ""):
        if blk.kind.startswith("h"):
            level = int(blk.kind[1])
            story.append(Paragraph(_inline_to_rl(blk.text), head[level]))
        elif blk.kind == "p":
            story.append(Paragraph(_inline_to_rl(blk.text), body))
        elif blk.kind == "bullet":
            story.append(Paragraph("&bull;&nbsp;&nbsp;" + _inline_to_rl(blk.text), body))
        elif blk.kind == "ordered":
            # We don't track numbering across blocks; render with a leading
            # bullet variant so it visually distinguishes from unordered
            # without misleading consecutive numbers.
            story.append(Paragraph("&#9702;&nbsp;&nbsp;" + _inline_to_rl(blk.text), body))
        elif blk.kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color="#888888"))

    doc.build(story)
    return buf.getvalue()


# --- DOCX --------------------------------------------------------------------

def render_docx(note: "Note", *, template_name: Optional[str] = None) -> bytes:
    document = Document()
    document.add_heading(_doc_title(note, template_name), level=0)

    if note.note_date:
        p = document.add_paragraph()
        p.add_run("Date: ").bold = True
        p.add_run(note.note_date.strftime("%B %d, %Y"))
    if note.author_name:
        p = document.add_paragraph()
        p.add_run("Author: ").bold = True
        p.add_run(note.author_name)
    if template_name:
        p = document.add_paragraph()
        p.add_run("Template: ").bold = True
        p.add_run(template_name)
    participants = _participant_names(note)
    if participants:
        p = document.add_paragraph()
        p.add_run("Participants: ").bold = True
        p.add_run(", ".join(participants))

    # Blank line + visual separation before the body.
    document.add_paragraph()

    markdown = relabel_speakers(note.note_content_markdown or "", note.speaker_labels)
    for blk in _parse_blocks(markdown or ""):
        if blk.kind.startswith("h"):
            level = int(blk.kind[1])
            # python-docx supports headings 1-9; we already cap at 4.
            document.add_heading(blk.text, level=level)
        elif blk.kind == "p":
            _add_inline_runs_docx(document.add_paragraph(), blk.text)
        elif blk.kind == "bullet":
            _add_inline_runs_docx(document.add_paragraph(style="List Bullet"), blk.text)
        elif blk.kind == "ordered":
            _add_inline_runs_docx(document.add_paragraph(style="List Number"), blk.text)
        elif blk.kind == "hr":
            # python-docx has no native horizontal rule. A paragraph of
            # em-dashes is the conventional ASCII substitute and keeps the
            # file editable in Word without surprises.
            document.add_paragraph("―" * 40)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
